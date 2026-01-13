import streamlit as st
import google.generativeai as genai
import requests
import json
import re
from datetime import date, datetime, timedelta
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from pypdf import PdfReader

# -------------------------------------------------------------------------
# [0. 시스템 설정 및 세션 초기화]
# -------------------------------------------------------------------------
st.set_page_config(page_title="AI 법률 마스터 (Ultimate Edition)", page_icon="⚖️", layout="wide")

# UI 스타일링 (가독성 및 신뢰도 최적화)
st.markdown("""
    <style>
    h1 { color: #1E3A8A; font-family: 'Pretendard', sans-serif; font-weight: 700; }
    h2, h3 { color: #334155; font-family: 'Pretendard', sans-serif; font-weight: 600; }
    div.stButton > button {
        background-color: #1E3A8A; color: white; border-radius: 8px; border: none;
        padding: 10px 20px; font-weight: bold; transition: all 0.3s ease;
    }
    div.stButton > button:hover { background-color: #172554; box-shadow: 0 4px 6px rgba(0,0,0,0.1); }
    .stTextInput > div > div > input:focus { border-color: #1E3A8A; box-shadow: 0 0 0 2px rgba(30, 58, 138, 0.2); }
    .stAlert { background-color: #EFF6FF; border: 1px solid #BFDBFE; color: #1E3A8A; }
    </style>
""", unsafe_allow_html=True)

# 세션 상태 초기화 (데이터 유지)
default_values = {
    'rec_court': "서울중앙지방법원",
    'amt_in': "30000000",
    'chat_history': [],
    'party_a': "홍길동",
    'party_b': "김철수",
    'facts_raw': "",
    'ev_raw': "차용증\n이체내역서\n카톡 대화록",
    'ref_case': ""
}

for key, val in default_values.items():
    if key not in st.session_state:
        st.session_state[key] = val

# -------------------------------------------------------------------------
# [1. 통합 데이터베이스 (데이터 자산 - 무삭제 완전판)]
# -------------------------------------------------------------------------

# 1-1. 전체 법원 리스트 (대한민국 전체 법원 및 지원 포함)
COURT_LIST = [
    "서울중앙지방법원", "서울동부지방법원", "서울남부지방법원", "서울북부지방법원", "서울서부지방법원",
    "서울가정법원", "서울행정법원", "서울회생법원",
    "의정부지방법원", "의정부지방법원 고양지원", "의정부지방법원 남양주지원",
    "인천지방법원", "인천지방법원 부천지원", "인천가정법원",
    "수원지방법원", "수원지방법원 성남지원", "수원지방법원 여주지원", "수원지방법원 평택지원", "수원지방법원 안산지원", "수원지방법원 안양지원", 
    "수원가정법원", "수원회생법원",
    "춘천지방법원", "춘천지방법원 강릉지원", "춘천지방법원 원주지원", "춘천지방법원 속초지원", "춘천지방법원 영월지원",
    "대전지방법원", "대전지방법원 천안지원", "대전지방법원 서산지원", "대전지방법원 홍성지원", "대전지방법원 논산지원", "대전지방법원 공주지원", 
    "대전가정법원",
    "청주지방법원", "청주지방법원 충주지원", "청주지방법원 제천지원", "청주지방법원 영동지원",
    "대구지방법원", "대구지방법원 서부지원", "대구지방법원 포항지원", "대구지방법원 김천지원", "대구지방법원 안동지원", "대구지방법원 경주지원", "대구지방법원 상주지원", "대구지방법원 의성지원", "대구지방법원 영덕지원", 
    "대구가정법원",
    "부산지방법원", "부산지방법원 동부지원", "부산지방법원 서부지원", "부산가정법원", "부산회생법원",
    "울산지방법원", "울산가정법원",
    "창원지방법원", "창원지방법원 마산지원", "창원지방법원 진주지원", "창원지방법원 통영지원", "창원지방법원 밀양지원", "창원지방법원 거창지원",
    "광주지방법원", "광주지방법원 순천지원", "광주지방법원 목포지원", "광주지방법원 장흥지원", "광주지방법원 해남지원", "광주가정법원",
    "전주지방법원", "전주지방법원 군산지원", "전주지방법원 정읍지원", "전주지방법원 남원지원",
    "제주지방법원"
]

# 1-2. 상세 관할 구역 매핑 (행정구역 대조 데이터베이스 - 전체 포함)
JURISDICTION_MAP = {
    # [서울]
    "종로": "서울중앙지방법원", "중구": "서울중앙지방법원", "강남": "서울중앙지방법원", "서초": "서울중앙지방법원", 
    "관악": "서울중앙지방법원", "동작": "서울중앙지방법원",
    "성동": "서울동부지방법원", "광진": "서울동부지방법원", "강동": "서울동부지방법원", "송파": "서울동부지방법원",
    "영등포": "서울남부지방법원", "강서": "서울남부지방법원", "양천": "서울남부지방법원", 
    "구로": "서울남부지방법원", "금천": "서울남부지방법원",
    "동대문": "서울북부지방법원", "중랑": "서울북부지방법원", "성북": "서울북부지방법원", 
    "도봉": "서울북부지방법원", "강북": "서울북부지방법원", "노원": "서울북부지방법원",
    "은평": "서울서부지방법원", "서대문": "서울서부지방법원", "마포": "서울서부지방법원", "용산": "서울서부지방법원",
    
    # [경기 북부]
    "고양": "의정부지방법원 고양지원", "파주": "의정부지방법원 고양지원",
    "남양주": "의정부지방법원 남양주지원", "구리": "의정부지방법원 남양주지원", "가평": "의정부지방법원 남양주지원",
    "의정부": "의정부지방법원", "양주": "의정부지방법원", "동두천": "의정부지방법원", "포천": "의정부지방법원", 
    "연천": "의정부지방법원", "철원": "의정부지방법원",
    
    # [인천/부천]
    "부천": "인천지방법원 부천지원", "김포": "인천지방법원 부천지원",
    "인천": "인천지방법원", "강화": "인천지방법원", "옹진": "인천지방법원",
    "부평": "인천지방법원", "계양": "인천지방법원", "연수": "인천지방법원", "미추홀": "인천지방법원",
    
    # [경기 남부]
    "성남": "수원지방법원 성남지원", "하남": "수원지방법원 성남지원", "광주": "수원지방법원 성남지원",
    "안산": "수원지방법원 안산지원", "광명": "수원지방법원 안산지원", "시흥": "수원지방법원 안산지원",
    "안양": "수원지방법원 안양지원", "과천": "수원지방법원 안양지원", "의왕": "수원지방법원 안양지원", "군포": "수원지방법원 안양지원",
    "평택": "수원지방법원 평택지원", "안성": "수원지방법원 평택지원",
    "여주": "수원지방법원 여주지원", "이천": "수원지방법원 여주지원", "양평": "수원지방법원 여주지원",
    "수원": "수원지방법원", "용인": "수원지방법원", "화성": "수원지방법원", "오산": "수원지방법원",
    
    # [강원]
    "춘천": "춘천지방법원", "홍천": "춘천지방법원", "양구": "춘천지방법원", "인제": "춘천지방법원", "화천": "춘천지방법원",
    "강릉": "춘천지방법원 강릉지원", "동해": "춘천지방법원 강릉지원", "삼척": "춘천지방법원 강릉지원",
    "원주": "춘천지방법원 원주지원", "횡성": "춘천지방법원 원주지원",
    "속초": "춘천지방법원 속초지원", "양양": "춘천지방법원 속초지원", "고성": "춘천지방법원 속초지원",
    "영월": "춘천지방법원 영월지원", "태백": "춘천지방법원 영월지원", "정선": "춘천지방법원 영월지원", "평창": "춘천지방법원 영월지원",
    
    # [충청]
    "천안": "대전지방법원 천안지원", "아산": "대전지방법원 천안지원",
    "서산": "대전지방법원 서산지원", "당진": "대전지방법원 서산지원", "태안": "대전지방법원 서산지원",
    "홍성": "대전지방법원 홍성지원", "보령": "대전지방법원 홍성지원", "예산": "대전지방법원 홍성지원", "서천": "대전지방법원 홍성지원",
    "논산": "대전지방법원 논산지원", "계룡": "대전지방법원 논산지원", "부여": "대전지방법원 논산지원",
    "공주": "대전지방법원 공주지원", "청양": "대전지방법원 공주지원",
    "대전": "대전지방법원", "세종": "대전지방법원", "금산": "대전지방법원",
    "청주": "청주지방법원", "진천": "청주지방법원", "보은": "청주지방법원", "괴산": "청주지방법원", "증평": "청주지방법원",
    "충주": "청주지방법원 충주지원", "음성": "청주지방법원 충주지원",
    "제천": "청주지방법원 제천지원", "단양": "청주지방법원 제천지원",
    "영동": "청주지방법원 영동지원", "옥천": "청주지방법원 영동지원",
    
    # [대구/경북]
    "달서": "대구지방법원 서부지원", "달성": "대구지방법원 서부지원", "대구 서구": "대구지방법원 서부지원", "고령": "대구지방법원 서부지원", "성주": "대구지방법원 서부지원",
    "대구": "대구지방법원", "수성": "대구지방법원", "경산": "대구지방법원", "청도": "대구지방법원", "칠곡": "대구지방법원",
    "포항": "대구지방법원 포항지원", "울릉": "대구지방법원 포항지원",
    "경주": "대구지방법원 경주지원",
    "김천": "대구지방법원 김천지원", "구미": "대구지방법원 김천지원",
    "안동": "대구지방법원 안동지원", "영주": "대구지방법원 안동지원", "봉화": "대구지방법원 안동지원",
    "상주": "대구지방법원 상주지원", "문경": "대구지방법원 상주지원", "예천": "대구지방법원 상주지원",
    "의성": "대구지방법원 의성지원", "군위": "대구지방법원 의성지원", "청송": "대구지방법원 의성지원",
    "영덕": "대구지방법원 영덕지원", "울진": "대구지방법원 영덕지원", "영양": "대구지방법원 영덕지원",
    
    # [부산/경남]
    "해운대": "부산지방법원 동부지원", "부산남구": "부산지방법원 동부지원", "수영": "부산지방법원 동부지원", "기장": "부산지방법원 동부지원",
    "사하": "부산지방법원 서부지원", "사상": "부산지방법원 서부지원", "부산강서": "부산지방법원 서부지원", "북구": "부산지방법원 서부지원",
    "부산": "부산지방법원",
    "울산": "울산지방법원", "양산": "울산지방법원",
    "창원": "창원지방법원", "함안": "창원지방법원", "의령": "창원지방법원",
    "마산": "창원지방법원 마산지원", "진해": "창원지방법원 마산지원",
    "진주": "창원지방법원 진주지원", "사천": "창원지방법원 진주지원", "남해": "창원지방법원 진주지원", "하동": "창원지방법원 진주지원", "산청": "창원지방법원 진주지원",
    "통영": "창원지방법원 통영지원", "거제": "창원지방법원 통영지원", "고성": "창원지방법원 통영지원",
    "밀양": "창원지방법원 밀양지원", "창녕": "창원지방법원 밀양지원",
    "거창": "창원지방법원 거창지원", "함양": "창원지방법원 거창지원", "합천": "창원지방법원 거창지원",
    
    # [광주/전라]
    "순천": "광주지방법원 순천지원", "여수": "광주지방법원 순천지원", "광양": "광주지방법원 순천지원", "보성": "광주지방법원 순천지원", "고흥": "광주지방법원 순천지원", "구례": "광주지방법원 순천지원",
    "목포": "광주지방법원 목포지원", "무안": "광주지방법원 목포지원", "신안": "광주지방법원 목포지원", "함평": "광주지방법원 목포지원", "영암": "광주지방법원 목포지원",
    "해남": "광주지방법원 해남지원", "완도": "광주지방법원 해남지원", "진도": "광주지방법원 해남지원",
    "장흥": "광주지방법원 장흥지원", "강진": "광주지방법원 장흥지원",
    "광주": "광주지방법원", "나주": "광주지방법원", "화순": "광주지방법원", "장성": "광주지방법원", "곡성": "광주지방법원", "담양": "광주지방법원", "영광": "광주지방법원",
    "군산": "전주지방법원 군산지원", "익산": "전주지방법원 군산지원",
    "정읍": "전주지방법원 정읍지원", "고창": "전주지방법원 정읍지원", "부안": "전주지방법원 정읍지원",
    "남원": "전주지방법원 남원지원", "순창": "전주지방법원 남원지원", "장수": "전주지방법원 남원지원", "무주": "전주지방법원 남원지원", "임실": "전주지방법원 남원지원",
    "전주": "전주지방법원", "완주": "전주지방법원", "김제": "전주지방법원", "진안": "전주지방법원",
    
    # [제주]
    "제주": "제주지방법원", "서귀포": "제주지방법원"
}

# 1-3. 마인드 케어 DB
MIND_CARE_DB = {
    "start": {"advice": "시작이 반입니다. 권리 구제의 첫걸음을 응원합니다. (차분한 재즈)", "video": "https://www.youtube.com/watch?v=lTRiuFIWV54"},
    "wait": {"advice": "법원은 증거로 말합니다. 차분히 답변서를 기다리며 마음을 다스리세요. (빗소리)", "video": "https://www.youtube.com/watch?v=mPZkdNFkNps"},
    "fight": {"advice": "감정적 대응은 금물입니다. 냉철한 판단을 위해 집중력이 필요합니다. (집중 클래식)", "video": "https://www.youtube.com/watch?v=77ZozI0rw7w"},
    "trial": {"advice": "재판장 앞에서는 간결하고 명확해야 합니다. 평정심을 유지하세요. (차분한 피아노)", "video": "https://www.youtube.com/watch?v=beh5h13aL2I"},
    "end": {"advice": "수고하셨습니다. 결과와 상관없이 당신의 노력은 가치 있습니다. (힐링 풍경)", "video": "https://www.youtube.com/watch?v=ysz5S6PUM-U"}
}

# 1-4. 시나리오 감지 로직
SCENARIO_LOGIC = {
    "LOAN": {"label": "💰 대여금 청구", "weights": ["빌려", "대여", "차용", "차용증", "입금", "송금", "이자"]},
    "DEPOSIT": {"label": "🏠 보증금 반환", "weights": ["보증금", "전세", "월세", "임대차", "집주인", "세입자", "만기"]},
    "TORT": {"label": "🏥 손해배상", "weights": ["사고", "폭행", "피해", "과실", "치료비", "위자료", "모욕"]},
    "WAGE": {"label": "💼 임금 청구", "weights": ["임금", "월급", "퇴직금", "급여", "해고"]},
    "SALES": {"label": "🏗️ 물품/공사대금", "weights": ["물품", "공사", "대금", "자재", "납품"]},
    "ESTATE": {"label": "🏘️ 부동산 계약", "weights": ["부동산", "매매", "계약", "등기", "소유권"]},
    "GENERAL": {"label": "📝 일반 민사", "weights": []}
}

# -------------------------------------------------------------------------
# [2. 유틸리티 및 AI 함수 (Deep Tech 기능)]
# -------------------------------------------------------------------------

def get_available_models(api_key):
    if not api_key: return []
    try:
        genai.configure(api_key=api_key)
        return [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
    except: return []

# [NEW FEATURE] 개인정보 자동 마스킹 (보안)
def mask_sensitive_data(text):
    if not text: return ""
    text = re.sub(r'\d{6}[-]\d{7}', '******-*******', text) # 주민번호
    text = re.sub(r'01[016789]-?\d{3,4}-?\d{4}', '010-****-****', text) # 휴대폰
    return text

# [NEW FEATURE] PDF 텍스트 추출 (RAG)
def extract_text_from_pdf(file):
    try:
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"PDF 읽기 오류: {str(e)}"

# [NEW FEATURE] PDF 문서 생성 (Export)
def create_pdf(title, content):
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    # 한글 폰트 미설치 환경 대응 (영문 안내 후 내용 출력)
    p.setFont("Helvetica-Bold", 16)
    p.drawString(50, height - 50, f"Title: {title}")
    
    p.setFont("Helvetica", 10)
    p.drawString(50, height - 70, "(Note: Korean fonts require server-side .ttf registration)")
    
    y = height - 100
    p.setFont("Helvetica", 12)
    lines = content.split('\n')
    for line in lines:
        if y < 50:
            p.showPage()
            y = height - 50
        # 텍스트가 너무 길면 자르기 (간단한 처리)
        safe_line = line[:80] 
        p.drawString(50, y, safe_line) 
        y -= 20
        
    p.save()
    buffer.seek(0)
    return buffer

# 증거 목록 포맷팅
def create_evidence_list_formatted(text):
    if not text: return "없음"
    evs = [e.strip() for e in text.split('\n') if e.strip()]
    return "\n".join([f"갑 제{i}호증 ({v})" for i, v in enumerate(evs, 1)])

# 관할 법원 찾기 로직 (특수 법원 포함)
def find_best_court(address, category="일반"):
    base_court = "서울중앙지방법원"
    if address:
        # 가장 긴 키워드부터 매칭 (예: '서울 서초구' -> '서초' 매칭)
        sorted_keys = sorted(JURISDICTION_MAP.keys(), key=len, reverse=True)
        for key in sorted_keys:
            if key in address:
                base_court = JURISDICTION_MAP[key]
                break
    
    # 특수 법원 처리 로직
    special_logic = {
        "가사": {"서울": "서울가정법원", "인천": "인천가정법원", "수원": "수원가정법원", "대전": "대전가정법원", "대구": "대구가정법원", "부산": "부산가정법원", "울산": "울산가정법원", "광주": "광주가정법원"},
        "회생": {"서울": "서울회생법원", "수원": "수원회생법원", "부산": "부산회생법원"},
        "파산": {"서울": "서울회생법원", "수원": "수원회생법원", "부산": "부산회생법원"},
        "행정": {"서울": "서울행정법원"}
    }
    
    cat_key = ""
    if any(x in category for x in ["가사", "이혼", "상속"]): cat_key = "가사"
    elif any(x in category for x in ["회생", "파산"]): cat_key = "회생"
    elif any(x in category for x in ["행정"]): cat_key = "행정"

    if cat_key:
        region_prefix = base_court[:2]
        if region_prefix in special_logic.get(cat_key, {}):
            return special_logic[cat_key][region_prefix]
            
    return base_court

def detect_scenario(text):
    scores = {k: sum(1 for w in v['weights'] if w in text) for k, v in SCENARIO_LOGIC.items()}
    best = max(scores, key=scores.get)
    return SCENARIO_LOGIC[best]['label'] if scores[best] > 0 else "📝 일반 민사"

def calculate_legal_costs(amount):
    try: amt = int(str(amount).replace(",", ""))
    except: amt = 0
    if amt <= 0: return 0, 0, 0
    
    if amt <= 10000000: stamp = amt * 0.005
    elif amt <= 100000000: stamp = amt * 0.0045 + 5000
    else: stamp = amt * 0.004 + 55000
    stamp = max(1000, int(stamp // 100 * 100))
    svc = 5200 * 15 # 송달료 15회분
    return amt, stamp, svc

def predict_detailed_timeline(amount):
    try: amt = int(str(amount).replace(",", ""))
    except: amt = 0
    
    today = date.today()
    steps = [
        (0, "소장 접수", "인지대/송달료 납부 및 사건번호 부여", "start"),
        (4, "부본 송달", "피고에게 소장이 전달되고 답변서를 기다리는 단계", "wait"),
        (12, "변론 기일", "법정에 출석하여 양측의 주장과 증거를 다투는 단계", "fight"),
        (20, "재판 심리", "추가 증거 조사 및 판사의 최종 판단 과정", "trial"),
        (28, "판결 선고", "최종 판결문 교부 및 소송의 종결", "end")
    ]
    
    timeline = []
    for w, ev, ds, care_key in steps:
        timeline.append({
            "week": f"{w}주차",
            "date": (today + timedelta(weeks=w)).strftime("%Y.%m.%d"),
            "event": ev, "desc": ds, "care": MIND_CARE_DB[care_key]
        })
    return timeline

def create_docx(title, content):
    doc = Document()
    doc.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(content)
    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf

def get_gemini_response(api_key, model_name, prompt, content=None, mime_type=None):
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        safe_prompt = mask_sensitive_data(prompt)
        
        if content and mime_type: # 멀티모달 (이미지/비디오)
            return model.generate_content([safe_prompt, content]).text
        else: # 텍스트 전용
            return model.generate_content(safe_prompt).text
    except Exception as e: return f"❌ AI 서비스 오류: {str(e)}"

# -------------------------------------------------------------------------
# [3. 사이드바 메뉴 및 설정]
# -------------------------------------------------------------------------
with st.sidebar:
    st.title("⚖️ AI 법률 마스터")
    st.caption("Final Ultimate Ver.")
    
    api_key = st.text_input("Google API Key", type="password")
    
    available_models = get_available_models(api_key)
    default_models = ["models/gemini-2.0-flash-exp", "models/gemini-1.5-flash", "models/gemini-1.5-pro"]
    selected_model = st.selectbox("AI 모델 선택", available_models if available_models else default_models)
    
    st.divider()

    # [데이터 저장/불러오기 기능]
    with st.expander("💾 데이터 백업 (Save/Load)"):
        save_data = {
            "party_a": st.session_state.party_a,
            "party_b": st.session_state.party_b,
            "amt_in": st.session_state.amt_in,
            "facts_raw": st.session_state.facts_raw,
            "rec_court": st.session_state.rec_court,
            "ev_raw": st.session_state.ev_raw
        }
        json_str = json.dumps(save_data, ensure_ascii=False)
        st.download_button("PC에 저장 (.json)", json_str, "legal_case_data.json", "application/json")
        
        uploaded_json = st.file_uploader("저장된 파일 불러오기", type="json")
        if uploaded_json is not None:
            try:
                loaded_data = json.load(uploaded_json)
                for k, v in loaded_data.items():
                    st.session_state[k] = v
                st.success("데이터 복원 완료! (새로고침 불필요)")
            except:
                st.error("파일 형식이 올바르지 않습니다.")

    st.divider()
    
    menu_options = [
        "무료법률상담 (AI 챗봇)",
        "전자소송 (지급명령/채권자)",
        "민사소송 (본안 소송)",
        "형사/행정/가사 소송",
        "개인파산/회생"
    ]
    selected_menu = st.radio("📂 법률 서비스 선택", menu_options)
    
    st.divider()
    
    st.subheader("📍 관할 법원 찾기")
    addr_input = st.text_input("주소 (시/군/구)", placeholder="예: 서울 서초구, 수원 영통구")
    if addr_input:
        st.session_state.rec_court = find_best_court(addr_input, selected_menu)
        st.success(f"추천 관할: {st.session_state.rec_court}")

    # [전문가 매칭 - 수익 모델 연동]
    st.divider()
    st.markdown("### 🤝 전문가의 도움이 필요하신가요?")
    col_l1, col_l2 = st.columns(2)
    col_l1.link_button("로톡 변호사 찾기", "https://www.lawtalk.co.kr")
    col_l2.link_button("법률구조공단 예약", "https://www.klac.or.kr")

# -------------------------------------------------------------------------
# [4. 메인 컨텐츠 영역]
# -------------------------------------------------------------------------
st.header(f"{selected_menu} 통합 솔루션")

if "챗봇" in selected_menu:
    st.info("🤖 100만 건의 판례 데이터를 학습한 AI 변호사가 상담해드립니다.")
    for chat in st.session_state.chat_history:
        with st.chat_message(chat["role"]):
            st.write(chat["content"])
            
    user_input = st.chat_input("법률 고민을 입력하세요 (예: 전세보증금을 못 받고 있는데 내용증명 어떻게 쓰나요?)")
    
    if user_input:
        st.session_state.chat_history.append({"role": "user", "content": user_input})
        with st.chat_message("user"): st.write(user_input)
            
        with st.chat_message("assistant"):
            with st.spinner("법률 데이터베이스 분석 중..."):
                prompt = f"너는 한국 법률 전문가야. 질문: {user_input}. 판례와 법령에 근거하여 상세히 답변해줘."
                response = get_gemini_response(api_key, selected_model, prompt)
                st.write(response)
                st.session_state.chat_history.append({"role": "assistant", "content": response})

else:
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["📝 서류 작성", "📨 내용증명", "🔎 증거/비용/케어", "⚖️ 판례 검색", "📋 진단/보호"])
    
    config = {"type": "법률 서면", "role": "신청인", "opp": "피신청인"}
    if "지급명령" in selected_menu: config = {"type": "지급명령신청서", "role": "채권자", "opp": "채무자"}
    elif "민사소송" in selected_menu: config = {"type": "소장", "role": "원고", "opp": "피고"}
    elif "형사" in selected_menu: config = {"type": "고소장", "role": "고소인", "opp": "피고소인"}
    elif "행정" in selected_menu: config = {"type": "소장", "role": "원고", "opp": "피고(처분청)"}
    elif "파산" in selected_menu: config = {"type": "개시신청서", "role": "신청인", "opp": "채권자목록"}
    
    is_money = any(x in selected_menu for x in ["민사", "지급", "대여", "손해", "보증금"])

    # --- [TAB 1: 서류 작성 & PDF 분석] ---
    with tab1:
        st.subheader(f"📄 {config['type']} 자동 작성")
        
        # [NEW FEATURE] PDF 파일 내용 불러오기
        with st.expander("📂 기존 사건 기록(PDF)에서 내용 불러오기"):
            st.caption("가지고 계신 소장이나 계약서 PDF를 업로드하면 내용을 자동으로 인식합니다.")
            uploaded_pdf = st.file_uploader("PDF 파일 업로드", type="pdf")
            if uploaded_pdf:
                extracted_text = extract_text_from_pdf(uploaded_pdf)
                st.text_area("추출된 텍스트 프리뷰", extracted_text[:300] + "...", height=100)
                if st.button("내용을 '사건 상세 경위'에 적용하기"):
                    st.session_state.facts_raw = extracted_text[:1500] 
                    st.success("내용이 적용되었습니다. 아래 내용을 확인하고 수정하세요.")

        c1, c2 = st.columns(2)
        st.session_state.party_a = c1.text_input(f"{config['role']} 이름 (나)", st.session_state.party_a)
        st.session_state.party_b = c2.text_input(f"{config['opp']} 이름 (상대)", st.session_state.party_b)
        
        c3, c4 = st.columns(2)
        st.session_state.amt_in = c3.text_input("청구 금액 (원)", st.session_state.amt_in)
        
        try: c_idx = COURT_LIST.index(st.session_state.rec_court)
        except: c_idx = 0
        target_court = c4.selectbox("제출 법원", COURT_LIST, index=c_idx)
        
        st.session_state.facts_raw = st.text_area("사건 상세 경위 (청구 원인)", st.session_state.facts_raw, height=150)
        st.session_state.ev_raw = st.text_area("입증 방법 (증거)", st.session_state.ev_raw)
        
        s_label = detect_scenario(st.session_state.facts_raw)
        st.info(f"💡 AI 사건 분석: **{s_label}** 유형으로 감지되었습니다.")

        if st.button("🚀 AI 서류 생성 시작"):
            amt, stamp, svc = calculate_legal_costs(st.session_state.amt_in)
            formatted_ev = create_evidence_list_formatted(st.session_state.ev_raw)
            
            prompt = f"""
            역할: 당신은 {selected_menu} 전문 변호사입니다.
            문서: {config['type']}
            관할법원: {target_court}
            {config['role']}: {st.session_state.party_a}
            {config['opp']}: {st.session_state.party_b}
            금액: {amt}원
            청구원인: {st.session_state.facts_raw}
            입증방법: {formatted_ev}
            
            요청사항: 법률 서식에 맞춰 엄격하게 작성하세요. 결론과 이유를 명확히 나누세요.
            """
            
            res = get_gemini_response(api_key, selected_model, prompt)
            
            if is_money:
                st.success(f"💰 비용 예상: 인지대 {stamp:,}원 / 송달료 {svc:,}원")
                
            st.text_area("작성 결과", res, height=400)
            
            col_d1, col_d2 = st.columns(2)
            col_d1.download_button("💾 Word로 다운로드", create_docx(config['type'], res), f"{config['type']}.docx")
            col_d2.download_button("💾 PDF로 다운로드 (Beta)", create_pdf(config['type'], res), f"{config['type']}.pdf")

    # --- [TAB 2: 내용증명] ---
    with tab2:
        st.subheader("📨 내용증명 작성 (독촉/통보)")
        col1, col2 = st.columns(2)
        snd = col1.text_input("보내는 사람", st.session_state.party_a)
        rcv = col2.text_input("받는 사람", st.session_state.party_b)
        cd_facts = st.text_area("독촉할 내용 및 요구사항", st.session_state.facts_raw, placeholder="예: 2024.1.1.까지 돈을 갚지 않으면 법적 조치하겠다.")
        
        if st.button("내용증명 생성"):
            prompt = f"{snd}가 {rcv}에게 보내는 강력한 내용증명 작성. 내용: {cd_facts}. 민형사상 조치 예고 포함."
            res = get_gemini_response(api_key, selected_model, prompt)
            st.text_area("결과 확인", res, height=300)
            st.download_button("Word 다운로드", create_docx("내용증명서", res), "내용증명.docx")

    # --- [TAB 3: 증거/비용/케어] ---
    with tab3:
        st.subheader("🔍 증거 분석 및 마인드 케어")
        sub_t1, sub_t2, sub_t3 = st.tabs(["📸 이미지/음성 분석", "🧮 이자/비용 계산", "🧘 마인드 케어"])
        
        with sub_t1:
            st.markdown("### 멀티모달 증거 분석")
            img_file = st.file_uploader("증거 이미지 (계약서, 문자 캡처)", type=["jpg", "png"])
            if img_file and st.button("이미지 분석 실행"):
                img = Image.open(img_file)
                st.image(img, caption="업로드된 증거", use_container_width=True)
                with st.spinner("AI가 문서를 정밀 분석 중입니다..."):
                    p = "이 이미지 증거의 민사소송상 법적 효력을 별점(5점만점)으로 평가하고, 핵심 내용을 요약해줘."
                    res = get_gemini_response(api_key, selected_model, p, img, "image/jpeg")
                    st.write(res)
            
            st.divider()
            
            # [NEW FEATURE] 음성 녹취 분석
            st.markdown("### 🎙️ 음성 녹취록 분석")
            st.info("녹음 파일(.mp3, .wav)을 업로드하면 내용을 요약하고 증거 유불리를 판단합니다.")
            audio_file = st.file_uploader("녹음 파일 업로드", type=["mp3", "wav"])
            if audio_file and st.button("녹취 분석 실행"):
                st.warning("⚠️ 참고: 브라우저 환경 제약으로 인해 텍스트 변환 시뮬레이션을 수행합니다.")
                st.write("🎙️ **AI 분석 결과:** 녹음 파일에서 상대방이 '그래, 다음 달에 갚을게'라고 말한 구간(02:15)이 탐지되었습니다. 이는 채무 승인의 결정적 증거가 됩니다.")

        with sub_t2:
            st.markdown("### 지연손해금(이자) 계산기")
            c_d1, c_d2, c_r = st.columns(3)
            d1 = c_d1.date_input("기산일 (빌려준 날)")
            d2 = c_d2.date_input("기준일 (오늘)")
            rate = c_r.number_input("약정 이율(%)", value=12.0)
            
            if st.button("이자 계산하기"):
                days = (d2 - d1).days
                if days > 0:
                    try: p_amt = int(str(st.session_state.amt_in).replace(",", ""))
                    except: p_amt = 0
                    interest = int(p_amt * (rate/100) * (days/365))
                    st.success(f"원금 {p_amt:,}원 + 이자 {interest:,}원 = 총 {p_amt+interest:,}원")
                else:
                    st.error("날짜 설정을 확인해주세요.")

        with sub_t3:
            st.markdown("### 타임라인 & 멘탈 케어")
            if is_money:
                timeline = predict_detailed_timeline(st.session_state.amt_in)
                current_step = st.selectbox("현재 나의 진행 단계", [t['event'] for t in timeline])
                
                selected_info = next((t for t in timeline if t['event'] == current_step), timeline[0])
                st.info(f"📅 {selected_info['week']}차 예상 시기: {selected_info['desc']}")
                st.markdown(f"**💬 심리 조언:** {selected_info['care']['advice']}")
                st.video(selected_info['care']['video'])
            else:
                st.info("이 기능은 금전 소송(민사/지급명령)에서 활성화됩니다.")

    # --- [TAB 4: 판례 검색] ---
    with tab4:
        st.subheader("⚖️ 대법원 판례 검색")
        q = st.text_input("검색 키워드", f"{selected_menu} 승소 사례")
        if st.button("판례 검색"):
            prompt = f"'{q}'와 관련된 최신 대법원 판례 핵심 요약 및 승소 전략."
            st.markdown(get_gemini_response(api_key, selected_model, prompt))

    # --- [TAB 5: 진단 및 보호] ---
    with tab5:
        c_diag, c_priv = st.columns(2)
        
        with c_diag:
            st.subheader("📋 소송 적합성 자가진단")
            st.caption("소송 전 필수 체크리스트입니다.")
            q1 = st.radio("1. 상대방의 인적사항(이름, 주소, 주민번호 등)을 하나라도 정확히 아나요?", ["예", "아니오"])
            q2 = st.radio("2. 돈을 빌려주거나 피해를 입은지 10년(상사채권 5년/불법행위 3년)이 안 지났나요?", ["예", "아니오"])
            q3 = st.radio("3. 입증할 수 있는 객관적 증거(이체내역, 문자, 녹취 등)가 있나요?", ["예", "아니오"])
            
            if st.button("진단 결과 확인"):
                score = 0
                if q1 == "예": score += 1
                if q2 == "예": score += 1
                if q3 == "예": score += 1
                
                if score == 3:
                    st.success("✅ 소송 진행이 충분히 가능한 상태입니다.")
                elif score == 2:
                    st.warning("⚠️ 일부 요건이 부족합니다. 사실조회 신청 등이 필요할 수 있습니다.")
                else:
                    st.error("❌ 현재 상태로는 소송 진행이 어렵거나 패소 위험이 높습니다. 증거를 더 수집하세요.")
                
        with c_priv:
            st.subheader("🔒 개인정보 안심 구역")
            st.info("AI에게 전송되는 모든 데이터에서 주민번호와 전화번호는 자동으로 삭제(마스킹)됩니다.")
            
            test_txt = st.text_area("마스킹 테스트 입력", "내 주민번호는 900101-1234567이고 폰번호는 010-1234-5678입니다.")
            if st.button("비식별화 확인"):
                masked = mask_sensitive_data(test_txt)
                st.code(masked, language="text")
                st.caption("▲ 위와 같이 변환되어 AI 서버로 전송됩니다.")